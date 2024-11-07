import datetime
from tkinter import PhotoImage
import importlib
from docx import Document
from docx.shared import Inches
import mainFrame
import tkinter as tk
from tkinter import ttk
import datetime as dt
import dbHandler
import pytz

class RegisterFrame():
    """ Initializes Sales Records Frame """
    def __init__(self, parent, user, isAdmin:bool, emp_id:int):
        self.parent = parent
        self.user = user
        self.isAdmin = isAdmin
        self.emp_id = emp_id
        self.db_params = {
            "database" : "postgres",
            "user" : "Celestin",
            "password" : "10august",
            "host" : "localhost",
            "port" : "5432"
        }

        self.updating = False
        self.today = dt.datetime.today()
        self.timezone = pytz.timezone("Europe/Bucharest")
        self.brand_logo = PhotoImage(file="./images/store.png")
        self.header_img = PhotoImage(file="./images/sales_header.png")
        self.passlabel_img = PhotoImage(file="./images/passlabel_img.png")
        self.userlabel_img = PhotoImage(file="./images/userlabel_img.png")
        self.logout_label_img = PhotoImage(file="./images/logout_label.png")
        self.or_label_img = PhotoImage(file="./images/or_label_img.png")
        self.title2_img = PhotoImage(file="./images/user.png")
        self.regimg = PhotoImage(file="./images/regimg.png")
        self.logout_img = PhotoImage(file="./images/logout.png")
        self.login_button_img = PhotoImage(file="./images/arrow.png")
        self.signup_img = PhotoImage(file="./images/signup_img.png")
        self.cpass_img = PhotoImage(file="./images/confirm.png")
        self.refresh_img = PhotoImage(file="./images/refresh.png")
        self.mail_img = PhotoImage(file="./images/email.png")
        self.back_img = PhotoImage(file="./images/back.png")
        self.regtable_img = PhotoImage(file="./images/table_img.png")
        self.saleid_img = PhotoImage(file="./images/saleid.png")
        self.selectprod_img = PhotoImage(file="./images/select_prod.png")
        self.cname_img = PhotoImage(file="./images/customer_name.png")
        self.country_img = PhotoImage(file="./images/country.png")
        self.address_img = PhotoImage(file="./images/address.png")
        self.phone_img = PhotoImage(file="./images/phone.png")
        self.zip_img = PhotoImage(file="./images/zip_code.png")
        self.add_img = PhotoImage(file="./images/add.png")
        self.remove_img = PhotoImage(file="./images/remove.png")
        self.findsale_img = PhotoImage(file="./images/find_sale.png")
        self.qty_img = PhotoImage(file="./images/qty.png")
        self.search_img = PhotoImage(file="./images/search.png")
        self.reset_img = PhotoImage(file="./images/undo.png")
        self.GUI()

    def GUI(self):
        self.bgcolor = "#ecebeb"
        self.frame = tk.Frame(self.parent, height=900, width=900, bg=self.bgcolor)
        self.frame.pack()
        self.header = tk.Label(self.parent, image=self.header_img, bg=self.bgcolor)
        self.header.place(x=230, y=20)
        self.select_label = tk.Label(self.parent, image=self.selectprod_img, bg=self.bgcolor)
        self.select_label.place(x=20, y=520)
        self.select_strvar = tk.StringVar()
        self.get_products()
        self.products_cboxlist = []
        self.products_list = self.get_products()
        self.style = ttk.Style()


        for i in self.products_list:
            self.products_cboxlist.append(i)
        self.select_cbox = ttk.Combobox(self.parent, width=25, values=self.products_cboxlist,
                                        font=("Arial", 15, "normal"))
        self.select_cbox.place(x=260, y=537)
        # self.select_cbox.current(0)
        self.prefix_list = ['+40','+93', '+54', '+61', '+43', '+32', '+55', '+56', '+86', '+61', '+61', '+57', '+53',
                            '+56', '+20', '+33', '+49', '+30', '+36', '+91', '+62', '+98', '+39', '+81', '+60', '+52',
                            '+95', '+31', '+64', '+47', '+92', '+51', '+63', '+48', '+65', '+27', '+82', '+34',
                            '+94', '+46', '+41', '+66', '+90', '+44', '+58', '+84']
        self.prefix_cbox = ttk.Combobox(self.parent, width=4, values=self.prefix_list)
        # self.prefix_cbox.current(0)
        self.prefix_cbox.place(x=258, y=635)

        self.cname_label = tk.Label(self.parent, image=self.cname_img, bg=self.bgcolor)
        self.cname_label.place(x=20, y=570)
        self.phone_label = tk.Label(self.parent, image=self.phone_img, bg=self.bgcolor)
        self.phone_label.place(x=25, y=615)
        self.country_label = tk.Label(self.parent, image=self.country_img, bg=self.bgcolor)
        self.country_label.place(x=30, y=670)
        self.address_label = tk.Label(self.parent, image=self.address_img, bg=self.bgcolor)
        self.address_label.place(x=30, y=770)
        self.zip_label = tk.Label(self.parent, image=self.zip_img, bg=self.bgcolor)
        self.zip_label.place(x=30, y=720)
        self.add_img = self.add_img.subsample(2, 2)
        self.remove_img = self.remove_img.subsample(2, 2)
        self.add_label = tk.Label(self.parent, image=self.add_img, bg=self.bgcolor)
        self.add_label.place(x=100, y=835)
        self.remove_label = tk.Label(self.parent, image=self.remove_img, bg=self.bgcolor)
        self.remove_label.place(x=405, y=835)
        self.findsale_label = tk.Label(self.parent, image=self.findsale_img, bg=self.bgcolor)
        self.findsale_label.place(x=560, y=585)
        self.quantity_label = tk.Label(self.parent, image=self.qty_img, bg=self.bgcolor)
        self.quantity_label.place(x=567, y=531)
        self.search_img = self.search_img.subsample(2, 2)
        self.search_label = tk.Label(self.parent, image=self.search_img, bg=self.bgcolor)
        self.search_label.place(x=660, y=770)
        self.refresh_img = self.refresh_img.subsample(2, 2)
        self.refresh_label = tk.Label(self.parent, image=self.refresh_img, bg=self.bgcolor)
        self.refresh_label.place(x=815, y=525)
        self.refresh_img.subsample(2, 2)
        self.reset_label = tk.Label(self.parent, image=self.reset_img, bg=self.bgcolor)
        self.reset_label.place(x=260, y=833)

        self.today_header = tk.Label(self.parent, text="", bg=self.bgcolor,
                                     foreground="black", font=("Arial", 14, "italic"))
        self.today_header.place(x=690, y=875)
        self.timenow_header = tk.Label(self.parent, text="", bg=self.bgcolor, foreground="black",
                                       font=("Arial", 14, "italic"))
        self.timenow_header.place(x=604, y=875)
        self.logout_img = self.logout_img.subsample(2, 2)
        self.logout = tk.Label(self.parent, image=self.logout_img, background=self.bgcolor)
        self.logout.place(x=785, y=55)
        self.logout_label_img = self.logout_label_img.subsample(2, 2)
        self.logout_label = ttk.Label(self.parent, image=self.logout_label_img, background=self.bgcolor)
        self.logout_label.place(x=765, y=10)

        self.reset_label.bind("<Button>", self.resetlabel_on_click)
        self.reset_label.bind("<Leave>", self.resetlabel_off_hover)
        self.reset_label.bind("<Enter>", self.resetlabel_on_hover)
        self.add_label.bind("<Enter>", self.addlabel_on_hover)
        self.add_label.bind("<Button>", self.addlabel_on_click)
        self.add_label.bind("<Leave>", self.addlabel_off_hover)
        self.remove_label.bind("<Enter>", self.removelabel_on_hover)
        self.remove_label.bind("<Button>", self.removelabel_on_click)
        self.remove_label.bind("<Leave>", self.removelabel_off_hover)
        self.search_label.bind("<Enter>", self.searchlabel_on_hover)
        self.search_label.bind("<Button>", self.searchlabel_on_click)
        self.search_label.bind("<Leave>", self.searchlabel_off_hover)
        self.logout.bind("<Enter>", self.logout_on_hover)
        self.logout.bind("<Button>", self.logout_on_click)
        self.logout.bind("<Leave>", self.logout_off_hover)
        self.refresh_label.bind("<Enter>", self.refresh_on_hover)
        self.refresh_label.bind("<Button>", self.refresh_on_click)
        self.refresh_label.bind("<Leave>", self.refresh_off_hover)

        self.cname_entry = tk.Entry(self.parent, width=25, bg="white", highlightthickness=0,
                                    font=("Arial", 16, "normal"), fg="black")
        self.cname_entry.place(x=260, y=583)
        self.phone_entry = tk.Entry(self.parent, width=19, bg="white", highlightthickness=0,
                                    font=("Arial", 16, "normal"), fg="black")
        self.phone_entry.place(x=318, y=633)
        self.zipcode_entry = tk.Entry(self.parent, width=25, bg="white", highlightthickness=0,
                                      font=("Arial", 16, "normal"), fg="black")
        self.zipcode_entry.place(x=260, y=725)
        self.find_entry = tk.Entry(self.parent, width=25, bg="white", highlightthickness=0,
                                   font=("Arial", 16, "normal"), fg="black")
        self.find_entry.place(x=578, y=710)

        self.countries_list =  ['Andorra', 'Albania', 'Austria', 'Åland Islands', 'Bosnia and Herzegovina', 'Belgium',
                            'Bulgaria', 'Belarus', 'Switzerland', 'Cyprus', 'Czech Republic', 'Germany', 'Denmark',
                            'Estonia', 'Spain', 'Finland', 'Faroe Islands', 'France', 'United Kingdom', 'Guernsey',
                            'Greece', 'Croatia', 'Hungary', 'Ireland', 'Isle of Man', 'Iceland', 'Italy', 'Jersey',
                            'Liechtenstein', 'Lithuania', 'Luxembourg', 'Latvia', 'Monaco', 'Moldova, Republic of',
                            'Macedonia, The Former Yugoslav Republic of', 'Malta', 'Netherlands', 'Norway', 'Poland',
                            'Portugal', 'Romania', 'Russian Federation', 'Sweden', 'Slovenia', 'Svalbard and Jan Mayen',
                            'Slovakia', 'San Marino', 'Ukraine', 'Holy See (Vatican City State)']

        self.countries_cbox = ttk.Combobox(self.parent, width=25, values=self.countries_list,
                                           font=("Arial", 15, "normal"))
        self.countries_cbox.place(x=260, y=680)
        # self.countries_cbox.current(40)
        self.findby_cbox = ttk.Combobox(self.parent, width=25, values=("ID", "Date", "EmpID", "Product Name", "Customer"))
        self.findby_cbox.place(x=575, y=660)
        self.findby_cbox.current(3)


        self.address_txtbox = tk.Text(self.parent, height=3, width=29, fg="black", bg="white", highlightthickness=0,
                               highlightbackground="white", font=("Arial", 14, "normal"))
        self.address_txtbox.place(x=259, y=770)
        self.q_spinbox = tk.Spinbox(self.parent, from_=0, to=100, width=6, repeatdelay=500, repeatinterval=100,
                             font=("Arial", 12, "normal"), bg="white", highlightthickness=0,
                                    highlightbackground="white", fg="black")
        self.q_spinbox.config(state="normal", cursor="hand2", bd=3, justify="center", wrap=True)
        self.q_spinbox.place(x=505, y=538)


        self.style = ttk.Style(self.parent)
        self.style.theme_use("clam")
        self.tree = ttk.Treeview(height=17, columns=11)
        self.tree.place(x=33, y=150)
        self.style.configure("Treeview", background="white", foreground="black", fieldbackground="white")
        self.tree["columns"] = ("ID", "Date", "EmpID", "Product Name", "Customer Name", "Country",
                                "Address", "Contact", "ZIP", "Quantity", "Price")
        self.tree.column("#0", width=0, minwidth=0)
        self.tree.column("ID", width=30, minwidth=30)
        self.tree.column("Date", width=80, minwidth=50)
        self.tree.column("EmpID", width=50, minwidth=30)
        self.tree.column("Product Name", width=130, minwidth=50)
        self.tree.column("Customer Name", width=100, minwidth=50)
        self.tree.column("Country", width=80, minwidth=80)
        self.tree.column("Address", width=120, minwidth=80)
        self.tree.column("Contact", width=90, minwidth=50)
        self.tree.column("ZIP", width=70, minwidth=30)
        self.tree.column("Quantity", width=40, minwidth=30)
        self.tree.column("Price", width=55, minwidth=35)

        self.tree.heading("#0", text="Label", anchor="w")
        self.tree.heading("ID", text="ID")
        self.tree.heading("Date", text="Date")
        self.tree.heading("EmpID", text="EmpID")
        self.tree.heading("Product Name", text="Product Name")
        self.tree.heading("Customer Name", text="Customer Name")
        self.tree.heading("Country", text="Country")
        self.tree.heading("Address", text="Address")
        self.tree.heading("Contact", text="Contact")
        self.tree.heading("ZIP", text="ZIP")
        self.tree.heading("Quantity", text="Qty")
        self.tree.heading("Price", text="Price")

        self.back_img = self.back_img.subsample(2, 2)
        self.back_label = tk.Label(self.parent, image=self.back_img, bg=self.bgcolor)
        self.back_label.place(x=20, y=20)
        self.back_label.bind("<Enter>", self.blabel_on_hover)
        self.back_label.bind("<Button>", self.blabel_on_click)
        self.back_label.bind("<Leave>", self.blabel_off_hover)
        self.tree.bind("<ButtonRelease>", self.tree_on_click)
        self.update_clock()
        self.load_salestree()
    def tree_on_click(self, event):
        self.clear_entries()
        self.iid = self.tree.focus()
        self.tree.selection_set(self.iid)
        self.updating = True
        self.item = self.tree.item(self.iid)
        print(f"item: {self.item["values"]}")
        self.cname = self.item["values"][4]
        self.country = self.item["values"][5]
        self.address = self.item["values"][6]
        self.contact = self.item["values"][7]
        self.prefix = str(self.contact)
        self.zip = self.item["values"][8]
        self.cname_entry.insert(0, self.cname)
        self.select_cbox.insert(0, self.item["values"][3])
        self.countries_cbox.insert(0, self.country)
        self.q_spinbox.insert(0, self.item["values"][9])
        self.phone_entry.insert(0, self.prefix[2:])
        self.address_txtbox.insert(1.0, self.address)
        self.prefix_cbox.insert(0, f"+{self.prefix[0:2]}")
        self.zipcode_entry.insert(0, self.zip)
    def on_event(self, event):
        self.login_button.configure(activebackground="#FFFFFF", highlightthickness=2, highlightbackground="#b4b4b6")

    def on_default(self, event):
        event.widget.configure(activebackground=self.bgcolor, highlightthickness=0, highlightbackground="#b4b4b6")

    def on_click(self, event):
        self.login_button.configure(highlightthickness=3, highlightbackground="#9d9d9e")
        self.on_default(event)

    def blabel_on_hover(self, event):
        event.widget.configure(highlightthickness=2, highlightbackground="#b4b4b6")

    def blabel_off_hover(self, event):
        event.widget.configure(activebackground=self.bgcolor, highlightthickness=0,
                               highlightbackground=self.bgcolor)

    def blabel_on_click(self, event):
        event.widget.configure(highlightthickness=3, highlightbackground="#9d9d9e")
        self.on_default(event)
        if self.isAdmin:
            self.mod = importlib.import_module("adminFrame")
            self.frame.destroy()
            self.app = self.mod.AdminFrame(self.parent, self.user, self.isAdmin, self.emp_id)
        else:
            self.mod = importlib.import_module("userFrame")
            self.frame.destroy()
            self.app = self.mod.UserFrame(self.parent, self.user, self.isAdmin, self.emp_id)

    def get_products(self):
        self.mod = importlib.import_module("dbHandler")
        self.products = self.mod.DBHandler.get_products(self)
        self.count = 0
        self.list = []
        for i in self.products:
            self.count += 1
            self.list.append(i[1])
        return self.list
    ##
    def create_invoice(self, name, product, unit, price, sale_id):
        self.document = Document()
        self.document.add_picture(image_path_or_stream="./images/store.png", width = Inches(1))
        self.document.add_heading("Invoice", 0)
        self.p1 = self.document.add_paragraph("Dear ")
        self.p1.add_run(name).bold = True
        self.p1.add_run(",")
        self.p2 = self.document.add_paragraph("Please find attached invoice for your recent purchase of ")
        self.p2.add_run(str(unit)).bold = True
        self.p2.add_run(" units of  ")
        self.p2.add_run(product).bold = True
        self.p2.add_run(".")
        [self.document.add_paragraph("") for _ in range(2)]
        self.table = self.document.add_table(rows=1, cols=4)
        self.hdr_cells = self.table.rows[0].cells
        self.hdr_cells[0].text = "Product Name"
        self.hdr_cells[1].text = "Units"
        self.hdr_cells[2].text = "Unit price"
        self.hdr_cells[3].text = "Total price"
        for i in range(4):
            self.hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        self.row_cells = self.table.add_row().cells
        self.row_cells[0].text = product
        self.row_cells[1].text = f"{unit}"
        self.row_cells[2].text = f"{round(price)} RON"
        self.row_cells[3].text = f"{int(unit) * int(price)}"

        [self.document.add_paragraph("") for _ in range(10)]
        self.document.add_paragraph("Thank you for your purchase and we hope to see you again soon !")
        self.document.add_paragraph("Sincerely,")
        self.document.add_paragraph("Demo Company Inc.")
        self.document.save(f"Invoices/{name.capitalize()}.docx")

    def update_clock(self):
        self.raw_ts = dt.datetime.now(self.timezone)
        self.date_now = self.raw_ts.strftime("%d %b %Y")
        self.time_now = self.raw_ts.strftime("%H:%M:%S %p")
        self.timenow_header.config(text=self.date_now)
        self.today_header.config(text=f" | {self.time_now}")
        self.today_header.after(1000, self.update_clock)



    def addlabel_off_hover(self, event):
        event.widget.configure(activebackground=self.bgcolor, highlightthickness=0,
                               highlightbackground=self.bgcolor)

    def addlabel_on_hover(self, event):
        self.add_label.configure(highlightthickness=2, highlightbackground="#b4b4b6")

    def get_timestamp(self):
        self.now = datetime.datetime.now()
        self.year = self.now.year
        self.month = self.now.month
        self.day = self.now.day
        self.hour = self.now.hour
        self.minute = self.now.minute
        self.second = self.now.second
        self.timestamp = f"{self.year}-{self.month}-{self.day} {self.hour}:{self.minute}:{self.second}"
        return self.timestamp

    def addlabel_on_click(self, event):
        self.get_timestamp()
        self.product_name = self.select_cbox.get()
        self.product_quantity = self.q_spinbox.get()
        self.cname = self.cname_entry.get()
        self.prefix = self.prefix_cbox.get()
        self.contact = self.phone_entry.get()
        self.country = self.countries_cbox.get()
        self.qty = self.q_spinbox.get()
        self.zip = self.zipcode_entry.get()
        self.address = self.address_txtbox.get(1.0, tk.END)
        if not self.updating:
            if self.product_name != "" and self.product_quantity != "" and self.cname != "" and self.prefix != "":
                if self.contact != "" and self.address != "" and self.zip != "" and self.country != "":
                    self.empid = dbHandler.DBHandler.get_empid(self, self.user)
                    self.add_label.configure(highlightthickness=3, highlightbackground="#9d9d9e")
                    self.addlabel_off_hover(event)
                    self.price = dbHandler.DBHandler.get_price(self, self.product_name)
                    self.price = self.price[0][0]
                    self.total = int(self.price) * int(self.qty)
                    self.phone = f"{self.prefix}{self.contact}"
                    print(f"timestamp is: {self.timestamp}")
                    self.timestamp = self.get_timestamp()
                    self.container = [self.timestamp, self.empid, self.product_name, self.cname, self.country,
                                                self.address, self.phone, self.zip, self.total, self.qty]
                    self.option = tk.messagebox.askquestion("Stock Manager", "You are going to create a new sale record !\n\n"
                                                                             f"Date : {self.timestamp}\n"       
                                                                             f"Product : {self.product_name}\n"
                                                                             f"Quantity: {self.qty}\n"
                                                                             f"Customer: {self.cname}\n"
                                                                             f"Phone : {self.phone}\n"
                                                                             f"Country : {self.country}\n"
                                                                             f"Address: {self.address}\n\n"
                                                                             f"TOTAL: {self.total}\n"
                                                                             "Are you sure?")
                    if self.option == "yes":
                        self.sale_id = dbHandler.DBHandler.add_entry(self, self.container)
                        self.load_salestree()
                        for child in self.tree.get_children():
                            self.iid = child
                            self.item = self.tree.item(self.iid)
                            print(f"qty: {self.qty} and price: {self.total}")
                            self.price_unit = (self.total / int(self.qty))
                        self.create_invoice(self.cname, self.product_name, self.qty, self.price_unit, self.sale_id)
                        self.updating = False
                        self.clear_entries()
                        self.load_salestree()
                        return tk.messagebox.showinfo("Inventory Manager", "You have succesfully created a new sale record !")
                    else: return
            else: return tk.messagebox.showerror("Stock Manager", "You forgot to fill an entry !")
        else:
            self.iid = self.tree.focus()
            self.item = self.tree.item(self.iid)
            self.id = self.item["values"][0]
            self.product_timestamp = self.item["values"][1]
            self.pname = self.select_cbox.get()
            self.quantity = self.q_spinbox.get()
            self.customer = self.cname_entry.get()
            self.pref = self.prefix_cbox.get()
            self.phone = self.phone_entry.get()

            self.countryvar = self.countries_cbox.get()
            self.zipp = self.zipcode_entry.get()
            self.addrss = self.address_txtbox.get(1.0, tk.END)
            self.product_price = dbHandler.DBHandler.get_price(self, self.pname)
            self.total = int(self.product_price[0][0]) * int(self.quantity)
            self.update = {
                "id" : self.id,
                "timestamp" : self.product_timestamp,
                "pname" : self.pname,
                "qty" : self.quantity,
                "cname" : self.customer,
                "prefix" : self.pref,
                "phone" : self.phone,
                "country" : self.countryvar,
                "zip" : self.zipp,
                "address" : self.addrss,
                "total" : self.total
            }
            self.option = tk.messagebox.askquestion("Stock Manager", f"You are going to update "
                                                                     f"the Sale ID:{self.id}\nAre you sure?")
            if self.option == "yes":
                self.response = dbHandler.DBHandler.update_sales(self, self.update)
                self.clear_entries()
                if self.response:
                    self.load_salestree()
                    return tk.messagebox.showinfo("Stock Manager", f"Sale ID:{self.id} succesfully updated !")
                else: return tk.messagebox.showerror("Stock Manager", f"Updating failed")

    def removelabel_off_hover(self, event):
        event.widget.configure(activebackground=self.bgcolor, highlightthickness=0,
                               highlightbackground=self.bgcolor)

    def removelabel_on_hover(self, event):
        self.remove_label.configure(highlightthickness=2, highlightbackground="#b4b4b6")

    def removelabel_on_click(self, event):
        self.remove_label.configure(highlightthickness=3, highlightbackground="#9d9d9e")
        self.removelabel_off_hover(event)
        if self.iid == 0: return
        self.iid = self.tree.focus()
        self.item = self.tree.item(self.iid)
        self.id = self.item["values"][0]
        self.option = tk.messagebox.askquestion("Stock Manager", f"You are going to delete Sale ID:{self.id}\n"
                                                                 f"Are you sure?")
        if self.option == "yes":
            self.response = dbHandler.DBHandler.remove_sale(self, self.id)
            if self.response:
                self.load_salestree()
                self.clear_entries()
                return tk.messagebox.showinfo("Stock Manager", f"You have deleted Sale ID:{self.id} !")

    def searchlabel_off_hover(self, event):
        event.widget.configure(activebackground=self.bgcolor, highlightthickness=0,
                               highlightbackground=self.bgcolor)

    def searchlabel_on_hover(self, event):
        self.search_label.configure(highlightthickness=2, highlightbackground="#b4b4b6")



    def searchlabel_on_click(self, event):
        self.search_label.configure(highlightthickness=3, highlightbackground="#9d9d9e")
        self.searchlabel_off_hover(event)
        self.method = self.findby_cbox.get()
        if self.method == "ID":
            self.method = "sale_id"
        elif self.method == "EmpID":
            self.method = "emp_id"
        elif self.method == "Date":
            self.method = "date"
        elif self.method == "Product Name":
            self.method = "product_name"
        elif self.method == "Customer":
            self.method = "customer_name"
        self.target = self.find_entry.get()
        self.response = dbHandler.DBHandler.search_sale(self, self.method, self.target)
        self.clear_tree()
        for i in self.response:
            self.tree.insert("", tk.END, values=(i[0], i[1], i[2], i[3], i[4], i[5], i[6], i[7], i[8],
                                                 i[9], i[10], i[1]))

    def clear_entries(self):
        for i in self.tree.selection():
            self.tree.selection_remove(i)
        self.updating = False
        self.select_cbox.delete(0, tk.END)
        self.q_spinbox.delete(0, tk.END)
        self.cname_entry.delete(0, tk.END)
        self.countries_cbox.delete(0, tk.END)
        self.phone_entry.delete(0, tk.END)
        self.address_txtbox.delete(1.0, tk.END)
        self.prefix_cbox.delete(0, tk.END)
        self.zipcode_entry.delete(0, tk.END)
        self.findby_cbox.delete(0, tk.END)
        self.find_entry.delete(0, tk.END)
        self.iid = 0

    def logout_off_hover(self, event):
        event.widget.configure(activebackground=self.bgcolor, highlightthickness=0,
                               highlightbackground=self.bgcolor)

    def logout_on_hover(self, event):
        self.logout.configure(highlightthickness=2, highlightbackground="#b4b4b6")

    def logout_on_click(self, event):
        self.logout.configure(highlightthickness=3, highlightbackground="#9d9d9e")
        self.logout_off_hover(event)
        self.mbox = tk.messagebox.askquestion("System", "You are going to be logged out !\n Are you sure?")
        if self.mbox == "yes":
            self.frame.destroy()
            self.app = mainFrame.MainFrame(self.parent)
        else: return

    def resetlabel_off_hover(self, event):
        event.widget.configure(activebackground=self.bgcolor, highlightthickness=0, highlightbackground=self.bgcolor)

    def resetlabel_on_click(self, event):
        self.reset_label.configure(highlightthickness=3, highlightbackground="#9d9d9e")
        self.resetlabel_off_hover(event)
        self.clear_entries()

    def resetlabel_on_hover(self, event):
        self.reset_label.configure(highlightthickness=2, highlightbackground="#b4b4b6")


    def refresh_off_hover(self, event):
        event.widget.configure(activebackground=self.bgcolor, highlightthickness=0, highlightbackground=self.bgcolor)

    def refresh_on_click(self, event):
        self.refresh_label.configure(highlightthickness=3, highlightbackground="#9d9d9e")
        self.refresh_off_hover(event)
        self.load_salestree()

    def refresh_on_hover(self, event):
        self.refresh_label.configure(highlightthickness=2, highlightbackground="#b4b4b6")

    def clear_tree(self):
        for item in self.tree.get_children():
            self.tree.delete(item)

    def load_salestree(self):
        self.record = self.mod.DBHandler.sales_record(self)
        self.clear_tree()
        self.count = 0
        for i in self.record:
            self.count += 1
            self.tree.insert("", tk.END, values=(i[0], i[1], i[2], i[3], i[4], i[5], i[6], i[7],i[8],i[9],i[10]))

